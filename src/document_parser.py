import logging
import os
import re
import tempfile
import zipfile
import importlib
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

try:
    from PyPDF2 import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    from docx import Document
    from docx.table import _Cell
except Exception:  # pragma: no cover
    Document = None
    _Cell = None

_BINARY_TEXT_RE = re.compile(rb"[ -~]{4,}")


class DocumentParser:
    """Extract text content from local files and uploaded bytes."""

    TEXT_EXTENSIONS = {
        ".txt",
        ".md",
        ".rst",
        ".json",
        ".csv",
        ".log",
        ".py",
        ".yaml",
        ".yml",
    }
    WORD_EXTENSIONS = {".docx", ".doc"}
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    def __init__(self, logger: Optional[logging.Logger] = None):
        self._logger = logger or logging.getLogger(__name__)
        self._markitdown = self._init_markitdown()
        self._marker_pdf_engine = self._init_marker_pdf_engine()
        self._pix2text_engine = self._init_pix2text_engine()

    @staticmethod
    def _preview_text(text: str, limit: int = 400) -> str:
        raw = str(text or "")
        if not raw:
            return ""
        compact = raw.replace("\r", "\\r").replace("\n", "\\n")
        if len(compact) <= limit:
            return compact
        return compact[:limit] + "..."

    def _log_extracted_text(self, source: str, text: str) -> None:
        preview = self._preview_text(text)
        self._logger.info(
            "DocumentParser 提取完成: source=%s length=%s preview=%s",
            source,
            len(str(text or "")),
            preview,
        )

    def _init_markitdown(self):
        try:
            module = importlib.import_module("markitdown")
            markitdown_cls = getattr(module, "MarkItDown", None)
        except Exception:
            markitdown_cls = None
        if markitdown_cls is None:
            self._logger.info("MarkItDown 不可用，回退到兼容解析器")
            return None
        try:
            return markitdown_cls()
        except Exception as exc:
            self._logger.warning("MarkItDown 初始化失败，回退到兼容解析器: %s", exc)
            return None

    def _init_marker_pdf_engine(self):
        try:
            converter_mod = importlib.import_module("marker.converters.pdf")
            models_mod = importlib.import_module("marker.models")
            converter_cls = getattr(converter_mod, "PdfConverter", None)
            create_model_dict = getattr(models_mod, "create_model_dict", None)
            if converter_cls is None:
                return None
            if callable(create_model_dict):
                return converter_cls(artifact_dict=create_model_dict())
            return converter_cls()
        except Exception:
            self._logger.info("PDF Marker 不可用，回退到当前 PDF 解析器")
            return None

    def _init_pix2text_engine(self):
        try:
            module = importlib.import_module("pix2text")
            cls = getattr(module, "Pix2Text", None)
            if cls is None:
                return None
            return cls()
        except Exception:
            self._logger.info("Pix2Text 不可用，回退到当前图片解析器")
            return None

    @staticmethod
    def _pick_markdown_text(result: object) -> str:
        if isinstance(result, dict):
            for key in ("text", "markdown", "content", "text_content", "md"):
                val = result.get(key)
                if isinstance(val, str) and val.strip():
                    return val
        if isinstance(result, (list, tuple)):
            parts = [DocumentParser._pick_markdown_text(x) for x in result]
            parts = [p for p in parts if p.strip()]
            if parts:
                return "\n\n".join(parts)
        for attr in ("text_content", "markdown", "content", "text"):
            val = getattr(result, attr, None)
            if isinstance(val, str) and val.strip():
                return val
        if isinstance(result, str) and result.strip():
            return result
        return ""

    @staticmethod
    def _table_to_markdown(rows: list[list[str]]) -> str:
        clean_rows = [[str(c or "").replace("\n", " ").strip() for c in row] for row in rows]
        clean_rows = [row for row in clean_rows if any(cell for cell in row)]
        if not clean_rows:
            return ""

        col_count = max(len(r) for r in clean_rows)
        normalized = [r + [""] * (col_count - len(r)) for r in clean_rows]
        header = normalized[0]
        sep = ["---"] * col_count
        body = normalized[1:]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(sep) + " |",
        ]
        for row in body:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    @staticmethod
    def _safe_cell_text(cell: Any) -> str:
        return str(getattr(cell, "text", "") or "").replace("\n", " ").strip()

    @staticmethod
    def _cell_grid_span(cell: Any) -> int:
        try:
            tc_pr = cell._tc.tcPr
            if tc_pr is not None and tc_pr.gridSpan is not None and tc_pr.gridSpan.val is not None:
                return max(1, int(tc_pr.gridSpan.val))
        except Exception:
            pass
        return 1

    @staticmethod
    def _cell_v_merge(cell: Any) -> str | None:
        try:
            tc_pr = cell._tc.tcPr
            if tc_pr is None or tc_pr.vMerge is None:
                return None
            val = tc_pr.vMerge.val
            if val is None:
                return "continue"
            return str(val)
        except Exception:
            return None

    @staticmethod
    def _render_complex_table_as_markdown(table_data: list[list[dict]]) -> str:
        if not table_data:
            return ""

        expanded_rows: list[list[str]] = []
        max_cols = 0
        for row in table_data:
            cells: list[str] = []
            for cell in row:
                text = str(cell.get("text", "") or "").replace("|", "\\|")
                rs = max(1, int(cell.get("row_span", 1) or 1))
                cs = max(1, int(cell.get("col_span", 1) or 1))
                if rs > 1 or cs > 1:
                    text = f"{text} [R{rs}xC{cs}]".strip()
                cells.append(text)
                for _ in range(cs - 1):
                    cells.append("")
            max_cols = max(max_cols, len(cells))
            expanded_rows.append(cells)

        if max_cols <= 0:
            return ""

        normalized_rows: list[list[str]] = []
        for row in expanded_rows:
            normalized_rows.append(row + [""] * (max_cols - len(row)))

        lines: list[str] = []
        for row in normalized_rows:
            lines.append("| " + " | ".join(row) + " |")

        lines.insert(1, "|" + "|".join([" --- " for _ in range(max_cols)]) + "|")
        return "\n".join(lines)

    @staticmethod
    def _extract_complex_tables_from_xml(doc_path: Path) -> list[list[list[dict]]]:
        if not zipfile.is_zipfile(doc_path):
            return []

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        try:
            with zipfile.ZipFile(doc_path, "r") as zf:
                raw = zf.read("word/document.xml")
        except Exception:
            return []

        try:
            root = ET.fromstring(raw)
        except Exception:
            return []

        tables_out: list[list[list[dict]]] = []
        for tbl in root.findall(".//w:tbl", ns):
            rows_out = DocumentParser._parse_xml_table_element(tbl, ns)
            if rows_out:
                tables_out.append(rows_out)
        return tables_out

    @staticmethod
    def _parse_xml_table_element(tbl: ET.Element, ns: dict[str, str]) -> list[list[dict]]:
        rows_out: list[list[dict]] = []
        ongoing: dict[int, dict] = {}

        for tr in tbl.findall("./w:tr", ns):
            row_cells: list[dict] = []
            used_cols: set[int] = set()
            col_cursor = 0

            for tc in tr.findall("./w:tc", ns):
                while col_cursor in ongoing or col_cursor in used_cols:
                    col_cursor += 1

                tc_pr = tc.find("./w:tcPr", ns)
                grid_span = 1
                if tc_pr is not None:
                    grid = tc_pr.find("./w:gridSpan", ns)
                    if grid is not None:
                        try:
                            grid_span = max(1, int(grid.attrib.get(f'{{{ns["w"]}}}val', "1")))
                        except Exception:
                            grid_span = 1

                v_merge_state = None
                if tc_pr is not None:
                    vm = tc_pr.find("./w:vMerge", ns)
                    if vm is not None:
                        raw_vm = vm.attrib.get(f'{{{ns["w"]}}}val')
                        v_merge_state = (raw_vm or "continue").strip().lower()

                texts = [
                    str(t.text or "").strip()
                    for t in tc.findall(".//w:t", ns)
                    if str(t.text or "").strip()
                ]
                text = "".join(texts).strip()

                if v_merge_state in {"continue", ""}:
                    base_cell = ongoing.get(col_cursor)
                    if base_cell is not None:
                        base_cell["row_span"] = int(base_cell.get("row_span", 1)) + 1
                    for c in range(col_cursor, col_cursor + grid_span):
                        used_cols.add(c)
                        if base_cell is not None:
                            ongoing[c] = base_cell
                    col_cursor += grid_span
                    continue

                cell_obj = {
                    "text": text,
                    "row_span": 1,
                    "col_span": grid_span,
                }
                row_cells.append(cell_obj)
                for c in range(col_cursor, col_cursor + grid_span):
                    used_cols.add(c)
                    if v_merge_state == "restart":
                        ongoing[c] = cell_obj
                    elif c in ongoing:
                        del ongoing[c]
                col_cursor += grid_span

            ended_cols = [c for c in ongoing.keys() if c not in used_cols]
            for c in ended_cols:
                ongoing.pop(c, None)

            if row_cells:
                rows_out.append(row_cells)
        return rows_out

    @staticmethod
    def _extract_docx_blocks_in_order(doc_path: Path) -> list[dict[str, Any]]:
        if not zipfile.is_zipfile(doc_path):
            return []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        try:
            with zipfile.ZipFile(doc_path, "r") as zf:
                raw = zf.read("word/document.xml")
            root = ET.fromstring(raw)
        except Exception:
            return []

        body = root.find(".//w:body", ns)
        if body is None:
            return []

        blocks: list[dict[str, Any]] = []
        for child in list(body):
            tag = str(child.tag or "")
            if tag.endswith("}p"):
                texts = [
                    str(t.text or "").strip()
                    for t in child.findall(".//w:t", ns)
                    if str(t.text or "").strip()
                ]
                txt = "".join(texts).strip()
                if txt:
                    blocks.append({"type": "text", "content": txt})
            elif tag.endswith("}tbl"):
                table_data = DocumentParser._parse_xml_table_element(child, ns)
                if table_data:
                    blocks.append({"type": "table", "content": table_data})
        return blocks

    @staticmethod
    def _extract_complex_table_data(doc_path: Path, table_index: int) -> list[list[dict]]:
        if Document is None:
            return []
        try:
            doc = Document(str(doc_path))
        except Exception:
            return []
        if table_index < 0 or table_index >= len(doc.tables):
            return []

        table = doc.tables[table_index]
        row_count = len(table.rows)
        col_count = len(table.columns) if table.columns else max(1, max((len(r.cells) for r in table.rows), default=1))

        reconstructed: list[list[dict | None]] = [[None for _ in range(col_count)] for _ in range(row_count)]

        for row_idx, row in enumerate(table.rows):
            col_cursor = 0
            for cell in row.cells:
                while col_cursor < col_count and reconstructed[row_idx][col_cursor] is not None:
                    col_cursor += 1
                if col_cursor >= col_count:
                    break

                col_span = DocumentParser._cell_grid_span(cell)
                v_merge = DocumentParser._cell_v_merge(cell)
                cell_info = {
                    "text": DocumentParser._safe_cell_text(cell),
                    "col_span": col_span,
                    "v_merge": v_merge,
                    "is_merge_start": (v_merge in {"restart", "true"}),
                    "is_placeholder": False,
                }
                reconstructed[row_idx][col_cursor] = cell_info

                for offset in range(1, col_span):
                    pos = col_cursor + offset
                    if pos < col_count:
                        reconstructed[row_idx][pos] = {"is_placeholder": True}
                col_cursor += col_span

        for c_idx in range(col_count):
            r_idx = 0
            while r_idx < row_count:
                cell = reconstructed[r_idx][c_idx]
                if not cell or cell.get("is_placeholder"):
                    r_idx += 1
                    continue

                v_merge = str(cell.get("v_merge") or "").lower()
                if cell.get("is_merge_start") or v_merge == "restart":
                    end = r_idx + 1
                    while end < row_count:
                        nxt = reconstructed[end][c_idx]
                        if not nxt:
                            break
                        nxt_vm = str(nxt.get("v_merge") or "").lower()
                        if nxt_vm in {"continue", "", "none"}:
                            end += 1
                            continue
                        break
                    span = max(1, end - r_idx)
                    cell["row_span"] = span
                    for rr in range(r_idx + 1, end):
                        nxt = reconstructed[rr][c_idx]
                        if nxt:
                            nxt["is_placeholder"] = True
                    r_idx = end
                else:
                    cell["row_span"] = 1
                    r_idx += 1

        formatted: list[list[dict]] = []
        for row in reconstructed:
            out_row: list[dict] = []
            for cell in row:
                if not cell or cell.get("is_placeholder"):
                    continue
                out_row.append(
                    {
                        "text": str(cell.get("text", "") or ""),
                        "row_span": int(cell.get("row_span", 1) or 1),
                        "col_span": int(cell.get("col_span", 1) or 1),
                    }
                )
            if out_row:
                formatted.append(out_row)
        return formatted

    @staticmethod
    def _extract_docx_table_blocks(path: Path) -> list[str]:
        table_blocks: list[str] = []
        seen: set[str] = set()

        if Document is not None:
            try:
                doc = Document(str(path))
            except Exception:
                doc = None
            if doc is not None:
                for table in doc.tables:
                    rows: list[list[str]] = []
                    for row in table.rows:
                        rows.append([cell.text.strip() for cell in row.cells])
                    md = DocumentParser._table_to_markdown(rows)
                    key = md.strip()
                    if key and key not in seen:
                        seen.add(key)
                        table_blocks.append(key)

        for md in DocumentParser._extract_docx_tables_from_xml(path):
            key = md.strip()
            if key and key not in seen:
                seen.add(key)
                table_blocks.append(key)

        return table_blocks

    @staticmethod
    def _extract_docx_tables_from_xml(path: Path) -> list[str]:
        if not zipfile.is_zipfile(path):
            return []

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        xml_names = ["word/document.xml"]
        table_blocks: list[str] = []

        try:
            with zipfile.ZipFile(path, "r") as zf:
                for name in zf.namelist():
                    low = name.lower()
                    if low.startswith("word/header") and low.endswith(".xml"):
                        xml_names.append(name)
                    elif low.startswith("word/footer") and low.endswith(".xml"):
                        xml_names.append(name)

                for name in xml_names:
                    try:
                        raw = zf.read(name)
                    except Exception:
                        continue
                    try:
                        root = ET.fromstring(raw)
                    except Exception:
                        continue
                    for tbl in root.findall(".//w:tbl", ns):
                        rows: list[list[str]] = []
                        for tr in tbl.findall("./w:tr", ns):
                            row: list[str] = []
                            for tc in tr.findall("./w:tc", ns):
                                texts = [
                                    str(t.text or "").strip()
                                    for t in tc.findall(".//w:t", ns)
                                    if str(t.text or "").strip()
                                ]
                                row.append("".join(texts).strip())
                            if any(cell for cell in row):
                                rows.append(row)
                        md = DocumentParser._table_to_markdown(rows)
                        if md.strip():
                            table_blocks.append(md.strip())
        except Exception:
            return []

        return table_blocks

    @staticmethod
    def _extract_docx_tables_markdown(path: Path) -> str:
        return "\n\n".join(DocumentParser._extract_docx_table_blocks(path))

    @staticmethod
    def _merge_docx_table_markdown(path: Path, markdown_text: str) -> str:
        table_blocks = DocumentParser._extract_docx_table_blocks(path)
        if not table_blocks:
            return markdown_text

        base = (markdown_text or "").strip()
        table_md = "\n\n".join(table_blocks)
        if not base:
            return table_md
        return f"{base}\n\n## 表格原文抽取\n\n{table_md}"

    def convert_file_to_markdown(self, path: Path) -> str:
        if self._markitdown is None:
            return ""
        try:
            result = self._markitdown.convert(str(path))
            text = self._pick_markdown_text(result)
            if path.suffix.lower() == ".docx":
                text = self._merge_docx_table_markdown(path, text)
            if text.strip():
                return text
        except Exception as exc:
            self._logger.debug("MarkItDown 转换失败: %s, error=%s", path, exc)
        return ""

    @staticmethod
    def read_text(path: Path) -> str:
        for enc in ("utf-8", "gb18030", "latin-1"):
            try:
                return path.read_text(encoding=enc, errors="ignore")
            except Exception:
                continue
        return ""

    @staticmethod
    def extract_text_from_xml_bytes(data: bytes) -> str:
        try:
            root = ET.fromstring(data)
        except Exception:
            return ""

        vals = []
        for elem in root.iter():
            if elem.text:
                t = elem.text.strip()
                if t:
                    vals.append(t)
        return "\n".join(vals)

    def read_pdf(self, path: Path) -> str:
        text = self._read_pdf_with_marker(path)
        if text.strip():
            return text

        if PdfReader is None:
            return ""
        try:
            reader = PdfReader(str(path))
        except Exception:
            return ""
        lines = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text)
        return "\n".join(lines)

    def _read_pdf_with_marker(self, path: Path) -> str:
        engine = self._marker_pdf_engine
        if engine is None:
            return ""
        try:
            if hasattr(engine, "convert"):
                result = engine.convert(str(path))
            elif hasattr(engine, "__call__"):
                result = engine(str(path))
            else:
                return ""
            text = self._pick_markdown_text(result)
            if text.strip():
                return text
        except Exception as exc:
            self._logger.debug("PDF Marker 解析失败: %s, error=%s", path, exc)
        return ""

    def _convert_word_to_pdf(self, src_path: Path, out_pdf: Path) -> bool:
        self._logger.info(
            "Word->PDF 开始: source=%s target=%s",
            src_path,
            out_pdf,
        )
        # Prefer Word COM conversion on Windows when available.
        if os.name == "nt":
            pythoncom = None
            com_inited = False
            word_app = None
            try:
                self._logger.info("Word->PDF 尝试转换器=win32com")
                pythoncom = importlib.import_module("pythoncom")
                pythoncom.CoInitialize()
                com_inited = True
                win32 = importlib.import_module("win32com.client")
                word_app = win32.DispatchEx("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = 0
                doc = None
                try:
                    doc = word_app.Documents.Open(str(src_path), ReadOnly=True)
                    # 17 == wdFormatPDF
                    doc.SaveAs(str(out_pdf), FileFormat=17)
                finally:
                    if doc is not None:
                        try:
                            doc.Close(False)
                        except Exception:
                            pass
                    if word_app is not None:
                        try:
                            word_app.Quit()
                        except Exception:
                            pass
                if out_pdf.exists() and out_pdf.stat().st_size > 0:
                    self._logger.info(
                        "Word->PDF 成功: converter=win32com output=%s size=%s",
                        out_pdf,
                        out_pdf.stat().st_size,
                    )
                    return True
            except Exception as exc:
                    self._logger.warning("Word->PDF 失败: converter=win32com source=%s error=%s", src_path, exc)
            finally:
                if pythoncom is not None and com_inited:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        # Fallback: try docx2pdf if installed.
        try:
            self._logger.info("Word->PDF 尝试转换器=docx2pdf")
            mod = importlib.import_module("docx2pdf")
            convert = getattr(mod, "convert", None)
            if callable(convert):
                convert(str(src_path), str(out_pdf))
                if out_pdf.exists() and out_pdf.stat().st_size > 0:
                    self._logger.info(
                        "Word->PDF 成功: converter=docx2pdf output=%s size=%s",
                        out_pdf,
                        out_pdf.stat().st_size,
                    )
                    return True
        except Exception as exc:
            self._logger.warning("Word->PDF 失败: converter=docx2pdf source=%s error=%s", src_path, exc)

        # Last fallback: LibreOffice/soffice CLI conversion.
        try:
            import subprocess

            self._logger.info("Word->PDF 尝试转换器=soffice")
            out_dir = out_pdf.parent
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(out_dir),
                    str(src_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if result.returncode == 0 and out_pdf.exists() and out_pdf.stat().st_size > 0:
                self._logger.info(
                    "Word->PDF 成功: converter=soffice output=%s size=%s",
                    out_pdf,
                    out_pdf.stat().st_size,
                )
                return True
            self._logger.warning(
                "Word->PDF 失败: converter=soffice source=%s returncode=%s stderr=%s",
                src_path,
                result.returncode,
                (result.stderr or "").strip()[:300],
            )
        except Exception as exc:
            self._logger.warning("Word->PDF 失败: converter=soffice source=%s error=%s", src_path, exc)

        self._logger.warning("Word->PDF 所有转换器均失败: source=%s", src_path)
        return False

    def _read_word_via_pdf_marker(self, path: Path) -> str:
        self._logger.info("Word->PDF->Marker 流程开始: source=%s", path)
        with tempfile.TemporaryDirectory(prefix="word_to_pdf_") as tmp_dir:
            pdf_path = Path(tmp_dir) / (path.stem + ".pdf")
            ok = self._convert_word_to_pdf(path, pdf_path)
            if not ok:
                self._logger.warning("Word->PDF->Marker 中止: PDF 转换失败 source=%s", path)
                return ""
            self._logger.info("Word->PDF->Marker 读取转换后 PDF: %s", pdf_path)
            text = self._read_pdf_with_marker(pdf_path)
            if text.strip():
                self._logger.info(
                    "Word->PDF->Marker 成功: source=%s extracted_length=%s",
                    path,
                    len(text),
                )
                return text
            self._logger.warning("Word->PDF->Marker 失败: Marker 返回空内容 source=%s", path)
        return ""

    def read_docx(self, path: Path) -> str:
        if Document is None:
            return ""
        blocks_in_order = self._extract_docx_blocks_in_order(path)
        if blocks_in_order:
            lines: list[str] = []
            table_idx = 0
            for block in blocks_in_order:
                b_type = str(block.get("type") or "")
                if b_type == "text":
                    txt = str(block.get("content") or "").strip()
                    if txt:
                        lines.append(txt)
                    continue
                if b_type == "table":
                    table_data = block.get("content") or []
                    md = self._render_complex_table_as_markdown(table_data)
                    if md.strip():
                        table_idx += 1
                        lines.append(f"### 表格 {table_idx}")
                        lines.append(md)
            if lines:
                return "\n\n".join(lines)

        try:
            doc = Document(str(path))
        except Exception:
            return ""
        lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

        complex_tables: list[str] = []
        for idx in range(len(doc.tables)):
            data = self._extract_complex_table_data(path, idx)
            if not data:
                continue
            md = self._render_complex_table_as_markdown(data)
            if md.strip():
                complex_tables.append(f"### 表格 {idx + 1}\n\n{md}")

        if complex_tables:
            lines.append("## 表格")
            lines.extend(complex_tables)
        else:
            table_md = self._extract_docx_tables_markdown(path)
            if table_md.strip():
                lines.append("## 表格")
                lines.append(table_md)
        return "\n\n".join(lines)

    def read_doc(self, path: Path) -> str:
        # Keep Word parsing on python-docx path as requested.
        text = self.read_docx(path)
        if text.strip():
            return text
        return self.read_binary_strings(path)

    def read_excel(self, path: Path) -> str:
        if pd is None:
            return self.read_binary_strings(path)
        try:
            sheets = pd.read_excel(path, sheet_name=None, header=None)
        except Exception:
            return self.read_binary_strings(path)

        tabulate_fn = None
        try:
            tabulate_mod = importlib.import_module("tabulate")
            tabulate_fn = getattr(tabulate_mod, "tabulate", None)
        except Exception:
            tabulate_fn = None

        rows = []
        for sheet_name, df in sheets.items():
            rows.append(f"# Sheet: {sheet_name}")
            data_rows = df.fillna("").astype(str).values.tolist()
            if not data_rows:
                continue
            if callable(tabulate_fn):
                try:
                    md_table = tabulate_fn(data_rows, headers="firstrow", tablefmt="pipe")
                    if str(md_table).strip():
                        rows.append(str(md_table))
                        continue
                except Exception:
                    pass
            for row in data_rows:
                vals = [x.strip() for x in row if x and x.strip()]
                if vals:
                    rows.append(" | ".join(vals))
        return "\n".join(rows)

    def read_image_table(self, path: Path) -> str:
        engine = self._pix2text_engine
        if engine is None:
            return ""
        try:
            for method_name in ("recognize", "recognize_text", "ocr"):
                method = getattr(engine, method_name, None)
                if not callable(method):
                    continue
                result = method(str(path))
                text = self._pick_markdown_text(result)
                if text.strip():
                    return text
        except Exception as exc:
            self._logger.debug("Pix2Text 识别失败: %s, error=%s", path, exc)
        return ""

    def read_ofd(self, path: Path) -> str:
        if not zipfile.is_zipfile(path):
            return self.read_binary_strings(path)

        texts = []
        try:
            with zipfile.ZipFile(path, "r") as zf:
                for name in zf.namelist():
                    lower = name.lower()
                    if not lower.endswith(".xml"):
                        continue
                    try:
                        data = zf.read(name)
                    except Exception:
                        continue
                    txt = self.extract_text_from_xml_bytes(data)
                    if txt.strip():
                        texts.append(txt)
        except Exception:
            return self.read_binary_strings(path)
        return "\n".join(texts)

    def read_pws(self, path: Path) -> str:
        if zipfile.is_zipfile(path):
            return self.read_ofd(path)
        return self.read_binary_strings(path)

    @staticmethod
    def read_binary_strings(path: Path) -> str:
        try:
            raw = path.read_bytes()
        except Exception:
            return ""

        candidates = []
        for enc in ("utf-8", "gb18030", "utf-16le", "latin-1"):
            try:
                s = raw.decode(enc, errors="ignore")
                if s:
                    candidates.append(s)
            except Exception:
                continue

        ascii_strings = [m.decode("ascii", errors="ignore") for m in _BINARY_TEXT_RE.findall(raw)]
        if ascii_strings:
            candidates.append("\n".join(ascii_strings))

        if not candidates:
            return ""
        candidates.sort(key=lambda x: len(x), reverse=True)
        return candidates[0]

    def extract_text(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext in self.WORD_EXTENSIONS:
            text = self.read_docx(path) if ext == ".docx" else self.read_doc(path)
            self._log_extracted_text(f"file:{path}", text)
            return text

        if path.suffix.lower() in self.IMAGE_EXTENSIONS:
            image_text = self.read_image_table(path)
            if image_text.strip():
                self._log_extracted_text(f"file:{path}", image_text)
                return image_text

        md_text = self.convert_file_to_markdown(path)
        if md_text.strip():
            self._log_extracted_text(f"file:{path}", md_text)
            return md_text

        text = self._extract_text_fallback(path)
        self._log_extracted_text(f"file:{path}", text)
        return text

    def _extract_text_fallback(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext in self.TEXT_EXTENSIONS:
            return self.read_text(path)
        if ext in self.WORD_EXTENSIONS:
            return self.read_docx(path) if ext == ".docx" else self.read_doc(path)
        if ext in self.IMAGE_EXTENSIONS:
            return self.read_image_table(path)
        if ext == ".pdf":
            return self.read_pdf(path)
        if ext == ".ofd":
            return self.read_ofd(path)
        if ext in {".xls", ".xlsx", ".xlsm"}:
            return self.read_excel(path)
        if ext in {".wps"}:
            return self.read_doc(path)
        if ext in {".pws"}:
            return self.read_pws(path)
        return self.read_binary_strings(path)

    @staticmethod
    def decode_uploaded_text_bytes(data: bytes, encoding: str | None = None) -> str:
        if encoding:
            try:
                return data.decode(encoding)
            except Exception as exc:
                raise ValueError(f"使用指定编码解码文件失败: {encoding}") from exc
        for enc in ("utf-8-sig", "utf-8", "gb18030", "utf-16", "utf-16le", "utf-16be"):
            try:
                return data.decode(enc)
            except Exception:
                continue
        return data.decode("utf-8", errors="replace")

    def extract_uploaded_file_text(
        self,
        filename: str,
        content: bytes,
        encoding: str | None = None,
        temp_dir: Path | None = None,
    ) -> str:
        data = bytes(content or b"")
        if not data:
            return ""

        ext = Path(filename).suffix.lower()
        if encoding is not None or ext in self.TEXT_EXTENSIONS:
            text = self.decode_uploaded_text_bytes(data, encoding=encoding)
            self._log_extracted_text(f"upload:{filename}", text)
            return text

        text = ""
        tmp_path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(
                mode="wb",
                suffix=(ext or ".bin"),
                dir=str(temp_dir) if temp_dir is not None else None,
                delete=False,
            ) as fp:
                fp.write(data)
                tmp_path = Path(fp.name)
            md_text = self.convert_file_to_markdown(tmp_path)
            if md_text.strip():
                self._log_extracted_text(f"upload:{filename}", md_text)
                return md_text
            text = self._extract_text_fallback(tmp_path)
        finally:
            if tmp_path is not None:
                try:
                    tmp_path.unlink()
                except Exception as exc:
                    self._logger.debug("清理上传临时文件失败: %s", exc)

        if not (text or "").strip():
            text = self.decode_uploaded_text_bytes(data, encoding=encoding)
        self._log_extracted_text(f"upload:{filename}", text)
        return text
