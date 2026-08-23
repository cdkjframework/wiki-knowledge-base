import logging
import os
import re
import subprocess
import sys
import zipfile
from contextlib import contextmanager
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from markitdown import MarkItDown
except Exception:  # pragma: no cover
    MarkItDown = None

from .ocr_parser import create_ocr_parser


_BINARY_TEXT_RE = re.compile(rb"[ -~]{4,}")
_CJK_RE = re.compile(r"[\u4e00-\u9fff]")


class BaseDocumentParser:
    _warned_once: set[str] = set()
    TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".json", ".csv", ".log", ".py", ".yaml", ".yml"}
    WORD_EXTENSIONS = {".docx", ".doc"}
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    @staticmethod
    def _detect_cuda_available() -> bool:
        try:
            import torch

            return bool(torch.cuda.is_available())
        except Exception:
            return False

    def __init__(
        self,
        logger: Optional[logging.Logger] = None,
        model_cache_dir: Optional[Path] = None,
        ocr_config: Optional[dict] = None,
    ):
        self._logger = logger or logging.getLogger(__name__)
        if ocr_config is None:
            import json

            config_path = Path("config.json")
            if config_path.exists():
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                ocr_config = config.get("knowledge_base", {}).get("ocr", {})
            else:
                ocr_config = {}
        default_cache_dir = Path("models/hf_cache").resolve()
        self._model_cache_dir = Path(model_cache_dir).resolve() if model_cache_dir is not None else default_cache_dir
        self._apply_model_cache_env()
        self._markitdown = self._init_markitdown()
        self._marker_pdf_engine = None
        self._pix2text_engine = None
        self._qwen_ocr_engine = None
        self._marker_pdf_checked = False
        self._pix2text_checked = False
        self._qwen_ocr_checked = False
        self._ocr_config = ocr_config if isinstance(ocr_config, dict) else {}
        # Marker / Pix2Text 跟 OCR 一样，只给商业版（KB-16）
        try:
            from ..commercial.edition import ocr_allowed as _ocr_allowed
        except ImportError:
            try:
                from commercial.edition import ocr_allowed as _ocr_allowed  # type: ignore
            except ImportError:
                try:
                    from src.commercial.edition import ocr_allowed as _ocr_allowed  # type: ignore
                except ImportError:
                    def _ocr_allowed() -> bool:  # type: ignore
                        return False

        edition_ocr_ok = bool(_ocr_allowed())
        self._marker_config_enabled = bool(self._ocr_config.get("enable_marker", True)) and edition_ocr_ok
        self._marker_cuda_available = self._detect_cuda_available()
        self._enable_marker = self._marker_config_enabled and self._marker_cuda_available
        self._enable_pix2text = bool(self._ocr_config.get("enable_pix2text", True)) and edition_ocr_ok
        if not edition_ocr_ok:
            self._logger.info("当前为社区版：已关闭 Marker / Pix2Text / OCR，仅保留文字层解析")
        llm_cfg = self._ocr_config.get("llm", {}) if isinstance(self._ocr_config.get("llm"), dict) else {}
        self._auto_install_missing_packages = bool(self._ocr_config.get("auto_install_missing_packages", True))
        self._ocr_model_name = str(llm_cfg.get("model_name") or self._ocr_config.get("model_name") or "Qwen/Qwen2-VL-2B-Instruct")
        self._ocr_prompt = str(llm_cfg.get("prompt") or self._ocr_config.get("prompt") or "请识别图片中的所有内容，并以 markdown 结构化文档返回。")
        self._ocr_device = str(llm_cfg.get("device") or self._ocr_config.get("device") or "cpu").strip().lower()
        self._ocr_dtype = str(llm_cfg.get("dtype") or self._ocr_config.get("dtype") or "auto").strip().lower()
        self._ocr_max_new_tokens = int(llm_cfg.get("max_new_tokens") or self._ocr_config.get("max_new_tokens") or 1024)
        self._ocr_min_pixels = int(llm_cfg.get("min_pixels") or self._ocr_config.get("min_pixels") or 0)
        self._ocr_max_pixels = int(llm_cfg.get("max_pixels") or self._ocr_config.get("max_pixels") or 0)
        self._pdf_ocr_dpi = int(self._ocr_config.get("pdf_ocr_dpi", 200))
        self._pdf_ocr_max_pages = int(self._ocr_config.get("pdf_ocr_max_pages", 0))
        # 商业实现在 commercial.business.ocr；社区拿到禁用桩
        self._ocr_parser = create_ocr_parser(
            self._ocr_config,
            model_cache_dir=self._model_cache_dir,
            logger=self._logger,
        )

    @staticmethod
    def _preview_text(text: str, limit: int = 400) -> str:
        raw = str(text or "")
        if not raw:
            return ""
        compact = raw.replace("\r", "\\r").replace("\n", "\\n")
        if len(compact) <= limit:
            return compact
        return compact[:limit] + "..."

    def _log_extracted_text(self, source: str, text: str) -> None:
        preview = self._preview_text(text)
        self._logger.info(
            "DocumentParser 提取完成: source=%s length=%s preview=%s",
            source,
            len(str(text or "")),
            preview,
        )

    @staticmethod
    def _hf_runtime_snapshot() -> dict[str, str]:
        return {
            "endpoint": str(os.getenv("HF_ENDPOINT") or os.getenv("HF_HUB_ENDPOINT") or "https://huggingface.co"),
            "cache_dir": str(os.getenv("HUGGINGFACE_HUB_CACHE") or os.getenv("HF_HOME") or ""),
            "hub_offline": str(os.getenv("HF_HUB_OFFLINE") or ""),
            "transformers_offline": str(os.getenv("TRANSFORMERS_OFFLINE") or ""),
            "datasets_offline": str(os.getenv("HF_DATASETS_OFFLINE") or ""),
        }

    def _log_hf_runtime(self, stage: str) -> None:
        snapshot = self._hf_runtime_snapshot()
        self._logger.info(
            "HF 运行环境: stage=%s endpoint=%s cache_dir=%s hub_offline=%s transformers_offline=%s datasets_offline=%s",
            stage,
            snapshot["endpoint"],
            snapshot["cache_dir"],
            snapshot["hub_offline"],
            snapshot["transformers_offline"],
            snapshot["datasets_offline"],
        )

    def _init_markitdown(self):
        if MarkItDown is None:
            self._warn_once("markitdown_missing", "info", "MarkItDown 不可用，回退到兼容解析器")
            return None
        try:
            return MarkItDown()
        except Exception as exc:
            self._warn_once("markitdown_init_failed", "warning", "MarkItDown 初始化失败，回退到兼容解析器: %s", exc)
            return None

    def _warn_once(self, key: str, level: str, message: str, *args: Any) -> None:
        if key in self._warned_once:
            return
        self._warned_once.add(key)
        getattr(self._logger, level, self._logger.warning)(message, *args)

    def _apply_model_cache_env(self) -> None:
        if self._model_cache_dir is None:
            return
        cache_dir = str(self._model_cache_dir)
        os.environ["HF_HOME"] = cache_dir
        os.environ["HUGGINGFACE_HUB_CACHE"] = cache_dir
        os.environ["TRANSFORMERS_CACHE"] = cache_dir
        os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
        os.environ.setdefault("HF_HUB_ENDPOINT", "https://hf-mirror.com")
        os.environ["MODEL_CACHE_DIR"] = str(self._model_cache_dir / "marker")
        os.environ["PIX2TEXT_HOME"] = str(self._model_cache_dir / "pix2text")
        os.environ["CNOCR_HOME"] = str(self._model_cache_dir / "cnocr")
        os.environ["CNSTD_HOME"] = str(self._model_cache_dir / "cnstd")
        os.environ["PADDLE_PDX_CACHE_HOME"] = str(self._model_cache_dir / "paddlex")

    def _ensure_optional_dependency(self, module_name: str, package_name: str) -> bool:
        try:
            self._import_optional_dependency(module_name)
            return True
        except Exception as exc:
            if not self._auto_install_missing_packages:
                self._logger.warning("Optional dependency missing: %s (%s). Auto-install disabled.", module_name, exc)
                return False
            self._logger.warning("Optional dependency missing: %s. Attempting pip install %s...", module_name, package_name)
            try:
                proc = subprocess.run(
                    [sys.executable, "-m", "pip", "install", package_name],
                    capture_output=True,
                    text=True,
                    check=False,
                )
                if proc.returncode != 0:
                    self._logger.warning(
                        "pip install %s failed (code=%s): %s",
                        package_name,
                        proc.returncode,
                        (proc.stderr or "").strip()[:300],
                    )
                    return False
                self._import_optional_dependency(module_name)
                return True
            except Exception as exc2:
                self._logger.warning("Auto-install failed for %s: %s", module_name, exc2)
                return False

    @staticmethod
    def _import_optional_dependency(module_name: str) -> None:
        if module_name == "fitz":
            import fitz  # noqa: F401

            return
        if module_name == "pdf2image":
            import pdf2image  # noqa: F401

            return
        if module_name == "markitdown":
            from markitdown import MarkItDown as _MarkItDown  # noqa: F401

            return
        raise ModuleNotFoundError(f"Unsupported optional dependency probe: {module_name}")

    @staticmethod
    @contextmanager
    def _hf_online_env(enabled: bool):
        if not enabled:
            yield
            return
        keys = ("HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "HF_DATASETS_OFFLINE")
        saved = {k: os.environ.get(k) for k in keys}
        for k in keys:
            os.environ.pop(k, None)
        try:
            try:
                import huggingface_hub.constants as hf_constants  # type: ignore

                hf_constants.HF_HUB_OFFLINE = False
            except Exception:
                pass
            try:
                import transformers.utils.hub as tf_hub_utils  # type: ignore

                if hasattr(tf_hub_utils, "_is_offline_mode"):
                    tf_hub_utils._is_offline_mode = False
            except Exception:
                pass
            yield
        finally:
            for k, v in saved.items():
                if v is None:
                    os.environ.pop(k, None)
                else:
                    os.environ[k] = v

    @staticmethod
    def _read_env_int(key: str, default: int, min_value: int | None = None, max_value: int | None = None) -> int:
        raw = os.getenv(key, "").strip()
        if raw:
            try:
                value = int(raw)
            except Exception:
                value = default
        else:
            value = default
        if min_value is not None:
            value = max(min_value, value)
        if max_value is not None:
            value = min(max_value, value)
        return value

    @staticmethod
    def _read_env_bool(key: str, default: bool) -> bool:
        raw = os.getenv(key, "").strip().lower()
        if not raw:
            return bool(default)
        return raw in {"1", "true", "yes", "on"}

    @staticmethod
    def _looks_like_mojibake(text: str, path: Path | None = None) -> bool:
        raw = str(text or "")
        if len(raw) < 40:
            return False
        nonspace = sum(1 for ch in raw if not ch.isspace())
        if nonspace <= 0:
            return False
        cjk_count = len(_CJK_RE.findall(raw))
        replacement = raw.count("\ufffd") + raw.count("�") + raw.count("□") + raw.count("■")
        replacement_ratio = replacement / nonspace if nonspace else 0.0
        if replacement_ratio > 0.02:
            return True
        if path is not None and _CJK_RE.search(path.name or ""):
            cjk_ratio = cjk_count / nonspace if nonspace else 0.0
            if cjk_ratio < 0.02:
                return True
        return False

    @staticmethod
    def _pick_markdown_text(result: object) -> str:
        if isinstance(result, dict):
            for key in ("text", "markdown", "content", "text_content", "md"):
                val = result.get(key)
                if isinstance(val, str) and val.strip():
                    return val
        if isinstance(result, (list, tuple)):
            parts = [BaseDocumentParser._pick_markdown_text(x) for x in result]
            parts = [p for p in parts if p.strip()]
            if parts:
                return "\n\n".join(parts)
        for attr in ("text_content", "markdown", "content", "text"):
            val = getattr(result, attr, None)
            if isinstance(val, str) and val.strip():
                return val
        if isinstance(result, str) and result.strip():
            return result
        return ""

    @staticmethod
    def _table_to_markdown(rows: list[list[str]]) -> str:
        clean_rows = [[str(c or "").replace("\n", " ").strip() for c in row] for row in rows]
        clean_rows = [row for row in clean_rows if any(cell for cell in row)]
        if not clean_rows:
            return ""
        col_count = max(len(r) for r in clean_rows)
        normalized = [r + [""] * (col_count - len(r)) for r in clean_rows]
        header = normalized[0]
        sep = ["---"] * col_count
        body = normalized[1:]
        lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
        for row in body:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    @staticmethod
    def _safe_cell_text(cell: Any) -> str:
        return str(getattr(cell, "text", "") or "").replace("\n", " ").strip()

    @staticmethod
    def _cell_grid_span(cell: Any) -> int:
        try:
            tc_pr = cell._tc.tcPr
            if tc_pr is not None and tc_pr.gridSpan is not None and tc_pr.gridSpan.val is not None:
                return max(1, int(tc_pr.gridSpan.val))
        except Exception:
            pass
        return 1

    @staticmethod
    def _cell_v_merge(cell: Any) -> str | None:
        try:
            tc_pr = cell._tc.tcPr
            if tc_pr is None or tc_pr.vMerge is None:
                return None
            val = tc_pr.vMerge.val
            if val is None:
                return "continue"
            return str(val)
        except Exception:
            return None

    @staticmethod
    def _render_complex_table_as_markdown(table_data: list[list[dict]]) -> str:
        if not table_data:
            return ""
        expanded_rows: list[list[str]] = []
        max_cols = 0
        for row in table_data:
            cells: list[str] = []
            for cell in row:
                text = str(cell.get("text", "") or "").replace("|", "\\|")
                rs = max(1, int(cell.get("row_span", 1) or 1))
                cs = max(1, int(cell.get("col_span", 1) or 1))
                if rs > 1 or cs > 1:
                    text = f"{text} [R{rs}xC{cs}]".strip()
                cells.append(text)
                for _ in range(cs - 1):
                    cells.append("")
            max_cols = max(max_cols, len(cells))
            expanded_rows.append(cells)
        if max_cols <= 0:
            return ""
        normalized_rows = [row + [""] * (max_cols - len(row)) for row in expanded_rows]
        lines = ["| " + " | ".join(row) + " |" for row in normalized_rows]
        lines.insert(1, "|" + "|".join([" --- " for _ in range(max_cols)]) + "|")
        return "\n".join(lines)

    @staticmethod
    def _extract_complex_tables_from_xml(doc_path: Path) -> list[list[list[dict]]]:
        if not zipfile.is_zipfile(doc_path):
            return []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        try:
            with zipfile.ZipFile(doc_path, "r") as zf:
                raw = zf.read("word/document.xml")
        except Exception:
            return []
        try:
            root = ET.fromstring(raw)
        except Exception:
            return []
        tables_out: list[list[list[dict]]] = []
        for tbl in root.findall(".//w:tbl", ns):
            rows_out = BaseDocumentParser._parse_xml_table_element(tbl, ns)
            if rows_out:
                tables_out.append(rows_out)
        return tables_out

    @staticmethod
    def _parse_xml_table_element(tbl: ET.Element, ns: dict[str, str]) -> list[list[dict]]:
        rows_out: list[list[dict]] = []
        ongoing: dict[int, dict] = {}
        for tr in tbl.findall("./w:tr", ns):
            row_cells: list[dict] = []
            used_cols: set[int] = set()
            col_cursor = 0
            for tc in tr.findall("./w:tc", ns):
                while col_cursor in ongoing or col_cursor in used_cols:
                    col_cursor += 1
                tc_pr = tc.find("./w:tcPr", ns)
                grid_span = 1
                if tc_pr is not None:
                    grid = tc_pr.find("./w:gridSpan", ns)
                    if grid is not None:
                        try:
                            grid_span = max(1, int(grid.attrib.get(f'{{{ns["w"]}}}val', "1")))
                        except Exception:
                            grid_span = 1
                v_merge_state = None
                if tc_pr is not None:
                    vm = tc_pr.find("./w:vMerge", ns)
                    if vm is not None:
                        raw_vm = vm.attrib.get(f'{{{ns["w"]}}}val')
                        v_merge_state = (raw_vm or "continue").strip().lower()
                texts = [str(t.text or "").strip() for t in tc.findall(".//w:t", ns) if str(t.text or "").strip()]
                text = "".join(texts).strip()
                if v_merge_state in {"continue", ""}:
                    base_cell = ongoing.get(col_cursor)
                    if base_cell is not None:
                        base_cell["row_span"] = int(base_cell.get("row_span", 1)) + 1
                    for c in range(col_cursor, col_cursor + grid_span):
                        used_cols.add(c)
                        if base_cell is not None:
                            ongoing[c] = base_cell
                    col_cursor += grid_span
                    continue
                cell_obj = {"text": text, "row_span": 1, "col_span": grid_span}
                row_cells.append(cell_obj)
                for c in range(col_cursor, col_cursor + grid_span):
                    used_cols.add(c)
                    if v_merge_state == "restart":
                        ongoing[c] = cell_obj
                    elif c in ongoing:
                        del ongoing[c]
                col_cursor += grid_span
            ended_cols = [c for c in ongoing.keys() if c not in used_cols]
            for c in ended_cols:
                ongoing.pop(c, None)
            if row_cells:
                rows_out.append(row_cells)
        return rows_out

    @staticmethod
    def _extract_docx_blocks_in_order(doc_path: Path) -> list[dict[str, Any]]:
        if not zipfile.is_zipfile(doc_path):
            return []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        try:
            with zipfile.ZipFile(doc_path, "r") as zf:
                raw = zf.read("word/document.xml")
            root = ET.fromstring(raw)
        except Exception:
            return []
        body = root.find(".//w:body", ns)
        if body is None:
            return []
        blocks: list[dict[str, Any]] = []
        for child in list(body):
            tag = str(child.tag or "")
            if tag.endswith("}p"):
                texts = [str(t.text or "").strip() for t in child.findall(".//w:t", ns) if str(t.text or "").strip()]
                txt = "".join(texts).strip()
                if txt:
                    blocks.append({"type": "text", "content": txt})
            elif tag.endswith("}tbl"):
                table_data = BaseDocumentParser._parse_xml_table_element(child, ns)
                if table_data:
                    blocks.append({"type": "table", "content": table_data})
        return blocks

    @staticmethod
    def _extract_complex_table_data(doc_path: Path, table_index: int) -> list[list[dict]]:
        if Document is None:
            return []
        try:
            doc = Document(str(doc_path))
        except Exception:
            return []
        if table_index < 0 or table_index >= len(doc.tables):
            return []
        table = doc.tables[table_index]
        row_count = len(table.rows)
        col_count = len(table.columns) if table.columns else max(1, max((len(r.cells) for r in table.rows), default=1))
        reconstructed: list[list[dict | None]] = [[None for _ in range(col_count)] for _ in range(row_count)]
        for row_idx, row in enumerate(table.rows):
            col_cursor = 0
            for cell in row.cells:
                while col_cursor < col_count and reconstructed[row_idx][col_cursor] is not None:
                    col_cursor += 1
                if col_cursor >= col_count:
                    break
                col_span = BaseDocumentParser._cell_grid_span(cell)
                v_merge = BaseDocumentParser._cell_v_merge(cell)
                cell_info = {
                    "text": BaseDocumentParser._safe_cell_text(cell),
                    "col_span": col_span,
                    "v_merge": v_merge,
                    "is_merge_start": (v_merge in {"restart", "true"}),
                    "is_placeholder": False,
                }
                reconstructed[row_idx][col_cursor] = cell_info
                for offset in range(1, col_span):
                    pos = col_cursor + offset
                    if pos < col_count:
                        reconstructed[row_idx][pos] = {"is_placeholder": True}
                col_cursor += col_span
        for c_idx in range(col_count):
            r_idx = 0
            while r_idx < row_count:
                cell = reconstructed[r_idx][c_idx]
                if not cell or cell.get("is_placeholder"):
                    r_idx += 1
                    continue
                v_merge = str(cell.get("v_merge") or "").lower()
                if cell.get("is_merge_start") or v_merge == "restart":
                    end = r_idx + 1
                    while end < row_count:
                        nxt = reconstructed[end][c_idx]
                        if not nxt:
                            break
                        nxt_vm = str(nxt.get("v_merge") or "").lower()
                        if nxt_vm in {"continue", "", "none"}:
                            end += 1
                            continue
                        break
                    span = max(1, end - r_idx)
                    cell["row_span"] = span
                    for rr in range(r_idx + 1, end):
                        nxt = reconstructed[rr][c_idx]
                        if nxt and not nxt.get("is_placeholder"):
                            nxt["is_placeholder"] = True
                    r_idx = end
                    continue
                cell["row_span"] = max(1, int(cell.get("row_span", 1) or 1))
                r_idx += 1
        rows_out: list[list[dict]] = []
        for row in reconstructed:
            row_cells: list[dict] = []
            for cell in row:
                if not cell or cell.get("is_placeholder"):
                    continue
                row_cells.append(
                    {
                        "text": str(cell.get("text", "") or ""),
                        "row_span": max(1, int(cell.get("row_span", 1) or 1)),
                        "col_span": max(1, int(cell.get("col_span", 1) or 1)),
                    }
                )
            if row_cells:
                rows_out.append(row_cells)
        return rows_out

    @staticmethod
    def _extract_docx_table_blocks(path: Path) -> list[str]:
        table_blocks: list[str] = []
        seen: set[str] = set()
        if Document is not None:
            try:
                doc = Document(str(path))
            except Exception:
                doc = None
            if doc is not None:
                for table in doc.tables:
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    md = BaseDocumentParser._table_to_markdown(rows)
                    key = md.strip()
                    if key and key not in seen:
                        seen.add(key)
                        table_blocks.append(key)
        for md in BaseDocumentParser._extract_docx_tables_from_xml(path):
            key = md.strip()
            if key and key not in seen:
                seen.add(key)
                table_blocks.append(key)
        return table_blocks

    @staticmethod
    def _extract_docx_tables_from_xml(path: Path) -> list[str]:
        if not zipfile.is_zipfile(path):
            return []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        xml_names = ["word/document.xml"]
        table_blocks: list[str] = []
        try:
            with zipfile.ZipFile(path, "r") as zf:
                for name in zf.namelist():
                    low = name.lower()
                    if low.startswith("word/header") and low.endswith(".xml"):
                        xml_names.append(name)
                    elif low.startswith("word/footer") and low.endswith(".xml"):
                        xml_names.append(name)
                for name in xml_names:
                    try:
                        raw = zf.read(name)
                    except Exception:
                        continue
                    try:
                        root = ET.fromstring(raw)
                    except Exception:
                        continue
                    for tbl in root.findall(".//w:tbl", ns):
                        rows: list[list[str]] = []
                        for tr in tbl.findall("./w:tr", ns):
                            row: list[str] = []
                            for tc in tr.findall("./w:tc", ns):
                                texts = [str(t.text or "").strip() for t in tc.findall(".//w:t", ns) if str(t.text or "").strip()]
                                row.append("".join(texts).strip())
                            if any(cell for cell in row):
                                rows.append(row)
                        md = BaseDocumentParser._table_to_markdown(rows)
                        if md.strip():
                            table_blocks.append(md.strip())
        except Exception:
            return []
        return table_blocks

    @staticmethod
    def _extract_docx_tables_markdown(path: Path) -> str:
        return "\n\n".join(BaseDocumentParser._extract_docx_table_blocks(path))

    @staticmethod
    def _merge_docx_table_markdown(path: Path, markdown_text: str) -> str:
        table_blocks = BaseDocumentParser._extract_docx_table_blocks(path)
        if not table_blocks:
            return markdown_text
        base = (markdown_text or "").strip()
        table_md = "\n\n".join(table_blocks)
        if not base:
            return table_md
        return f"{base}\n\n## 表格原文抽取\n\n{table_md}"

    @staticmethod
    def read_text(path: Path) -> str:
        for enc in ("utf-8", "gb18030", "latin-1"):
            try:
                return path.read_text(encoding=enc, errors="ignore")
            except Exception:
                continue
        return ""

    @staticmethod
    def extract_text_from_xml_bytes(data: bytes) -> str:
        try:
            root = ET.fromstring(data)
        except Exception:
            return ""
        vals = []
        for elem in root.iter():
            if elem.text:
                t = elem.text.strip()
                if t:
                    vals.append(t)
        return "\n".join(vals)

    @staticmethod
    def read_binary_strings(path: Path) -> str:
        try:
            raw = path.read_bytes()
        except Exception:
            return ""
        candidates = []
        for enc in ("utf-8", "gb18030", "utf-16le", "latin-1"):
            try:
                s = raw.decode(enc, errors="ignore")
                if s:
                    candidates.append(s)
            except Exception:
                continue
        ascii_strings = [m.decode("ascii", errors="ignore") for m in _BINARY_TEXT_RE.findall(raw)]
        if ascii_strings:
            candidates.append("\n".join(ascii_strings))
        if not candidates:
            return ""
        candidates.sort(key=lambda x: len(x), reverse=True)
        return candidates[0]

    @staticmethod
    def decode_uploaded_text_bytes(data: bytes, encoding: str | None = None) -> str:
        if encoding:
            try:
                return data.decode(encoding)
            except Exception as exc:
                raise ValueError(f"使用指定编码解码文件失败: {encoding}") from exc
        for enc in ("utf-8-sig", "utf-8", "gb18030", "utf-16", "utf-16le", "utf-16be"):
            try:
                return data.decode(enc)
            except Exception:
                continue
        return data.decode("utf-8", errors="replace")

    def _patch_optimum_onnxruntime_exports(self) -> None:
        return